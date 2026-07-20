from pathlib import Path

from ingestion.base import AbstractDocumentLoader
from ingestion.document import StructuredDocument, TextBlock
from ingestion.exceptions import IngestionError


class DocxLoader(AbstractDocumentLoader):
    def load(self, path: Path) -> StructuredDocument:
        try:
            from docx import Document
        except ImportError:
            raise IngestionError("python-docx is required to read .docx files.")

        try:
            doc = Document(str(path))
        except Exception as exc:
            raise IngestionError(f"Failed to open docx {path}: {exc}")

        blocks = []
        paragraphs = []

        for i, p in enumerate(doc.paragraphs):
            text = p.text
            if not text:
                continue
            paragraphs.append(text)
            blocks.append(
                TextBlock(
                    text=text,
                    page=0,
                    block_type="text",
                )
            )

        return StructuredDocument(
            text="\n".join(paragraphs),
            blocks=blocks,
            metadata={"paragraphs": len(paragraphs)},
        )
